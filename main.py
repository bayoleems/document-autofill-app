print("\U0001F1F3\U0001F1EC")

from flask import Flask, render_template, request, send_file, redirect, url_for
import uvicorn
from docx import Document
import os
import json
import re
import threading
import time

def download_and_redirect(file_path):
    time.sleep(1)
    os.remove(file_path)
    os.remove(fr"uploads/{file_path.split('filled_')[1]}")
    return "Follow me all social media platforms: @bayoleems"

def extract_words(text):
    pattern = r'\[\[(.*?)\]\]'
    matches = re.findall(pattern, text)
    return matches

def generate_html_form(words):
    html_content = '''
    <!DOCTYPE html>
    <html lang="en">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>Document Auto-Fill</title>
        <style>
            body {
                background-color: #222;
                color: #fff;
                font-family: Arial, sans-serif;
            }
            .form-container {
                max-width: 400px;
                margin: 50px auto;
                padding: 20px;
                border-radius: 10px;
                background-color: #333;
                box-shadow: 0 0 10px rgba(0, 0, 0, 0.5);
            }
            .form-group {
                margin-bottom: 20px;
            }
            label {
                display: block;
                margin-bottom: 5px;
            }
            input[type="text"] {
                width: 90%;
                padding: 10px;
                border-radius: 5px;
                border: 1px solid #666;
                background-color: #444;
                color: #fff;
            }
            input[type="submit"] {
                padding: 10px 20px;
                border-radius: 5px;
                border: none;
                background-color: #007bff;
                color: #fff;
                cursor: pointer;
            }
            input[type="submit"]:hover {
                background-color: #0056b3;
            }
        </style>
    </head>
    <body>
        <div class="form-container">
            <form action="/submit" method="post">
    '''

    for word in words:
        html_content += f'''
                <div class="form-group">
                    <label for="{word}">{word.capitalize()}</label>
                    <input type="text" id="{word}" name="{word}">
                </div>
        '''

    html_content += '''
                <input type="submit" value="Submit">
            </form>
        </div>
    </body>
    </html>
    '''

    return html_content


app = Flask(__name__)
app.secret_key = 'leems'

@app.route('/')
def index():
    return render_template('upload.html')

@app.route('/upload',  methods=['POST'])
def process():
    if 'file' not in request.files:

        return redirect(url_for('index'))

    file = request.files['file']

    if file.filename == '':

        return redirect(url_for('index'))

    if file:
        dir = 'uploads'

        # Check if the path exists
        if not os.path.exists(dir):
            # If the path does not exist, create it
            os.makedirs(dir)

        # Save the file to the upload directory
        filename = os.path.join(dir, file.filename)
        file.save(filename)

        with open('filename.json') as fn:
            info = json.load(fn)

        if info["file"] == "":
            info["file"] = filename 

            os.remove(r'filename.json')
            with open('filename.json', 'w') as fl:
                json.dump(info, fl)

        doc = Document(fr'{filename}')
        words = set()
        for paragraph in doc.paragraphs:
            for word in extract_words(paragraph.text):
                words.add(word)
        
                    
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for word in extract_words(paragraph.text):
                            words.add(word)

        return generate_html_form(words= list(words))
        
@app.route('/submit', methods=['POST'])
def submit():
    if request.method == 'POST':

        with open('filename.json') as fn:
            info = json.load(fn)
        filename = info["file"]

        os.remove(r'filename.json')
        with open('filename.json', 'w') as fl:
            info = {"file": ""}
            json.dump(info, fl)

        doc = Document(filename)
        data = {}
        for key, value in request.form.items():
            data[f'[[{key}]]'] = value

        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                for k,v in data.items():
                    if k in run.text:
                        font = run.font
                        size = font.size
                        run.text = run.text.replace(str(k), str(v))
                        run.font.size = size
                    
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            for k,v in data.items():
                                if k in run.text:
                                    font = run.font
                                    size = font.size
                                    run.text = run.text.replace(str(k), str(v))
                                    run.font.size = size

        name = f"filled_{filename.split('/')[-1]}"
        doc.save(name)
        threading.Thread(target=download_and_redirect, args=(name,)).start()
        return send_file(name, as_attachment=True)

if __name__ == '__main__':
    app.run(debug=True)
    uvicorn.run("main:app", host = "0.0.0.0", port = 5050, log_level = "info", reload = True)